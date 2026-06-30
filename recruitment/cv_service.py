from __future__ import annotations

import base64
import json
from typing import Dict, Optional, Tuple

import docx
import fitz

from .config import get_recruitment_settings
from .llm_client import get_llm_client

try:
    from docling.document_converter import DocumentConverter

    DOCLING_AVAILABLE = True
except Exception:
    DOCLING_AVAILABLE = False


def extract_text_with_pymupdf(file_path: str) -> Dict:
    try:
        document = fitz.open(file_path)
        pages = []
        full_text = []

        for index, page in enumerate(document):
            page_text = page.get_text("text")
            full_text.append(page_text)
            pages.append(
                {
                    "page_number": index + 1,
                    "text": page_text,
                    "text_length": len(page_text),
                }
            )

        result = {
            "text": "\n".join(full_text),
            "pages": pages,
            "metadata": document.metadata,
            "page_count": len(document),
        }
        document.close()
        return result
    except Exception:
        return {"text": "", "pages": [], "metadata": {}, "page_count": 0}


def extract_links_from_pdf(file_path: str) -> Dict[str, list]:
    try:
        document = fitz.open(file_path)
        linkedin = []
        github = []
        others = []

        for page in document:
            for link in page.get_links():
                uri = link.get("uri", "")
                if not uri:
                    continue

                uri_lower = uri.lower()
                if "linkedin.com" in uri_lower and uri not in linkedin:
                    linkedin.append(uri)
                elif "github.com" in uri_lower and uri not in github:
                    github.append(uri)
                elif uri not in others:
                    others.append(uri)

        document.close()
        return {"linkedin": linkedin, "github": github, "others": others}
    except Exception:
        return {"linkedin": [], "github": [], "others": []}


def extract_with_docling(file_path: str) -> Optional[Dict]:
    if not DOCLING_AVAILABLE:
        return None

    try:
        result = DocumentConverter().convert(file_path)
        return {
            "text": result.document.export_to_text(),
            "markdown": result.document.export_to_markdown(),
            "structure": str(result.document),
        }
    except Exception:
        return None


def extract_text_from_docx(file_path: str) -> str:
    try:
        document = docx.Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

        tables = []
        for table in document.tables:
            for row in table.rows:
                tables.append(" | ".join(cell.text.strip() for cell in row.cells))

        return "\n".join(paragraphs + tables)
    except Exception:
        return ""


def extract_text_from_image(image_path: str) -> str:
    with open(image_path, "rb") as image_file:
        b64_data = base64.b64encode(image_file.read()).decode()

    image_type = "image/png" if image_path.lower().endswith(".png") else "image/jpeg"

    settings = get_recruitment_settings()
    llm_client = get_llm_client()
    for model in [settings.ocr_model, "google/gemini-flash-1.5"]:
        try:
            response = llm_client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{image_type};base64,{b64_data}"},
                            },
                            {
                                "type": "text",
                                "text": "Extract all text from this CV/resume image. Preserve structure. Return only the text.",
                            },
                        ],
                    }
                ],
                max_tokens=3000,
            )
            return response.choices[0].message.content or ""
        except Exception:
            continue

    return ""


def process_file(file_path: str) -> Tuple[str, Dict]:
    extension = file_path.lower().split(".")[-1]
    links = {"linkedin": [], "github": [], "others": []}

    if extension == "pdf":
        text = extract_text_with_pymupdf(file_path).get("text", "")
        links = extract_links_from_pdf(file_path)
        return text, links

    if extension in ("docx", "doc"):
        return extract_text_from_docx(file_path), links

    if extension in ("jpg", "jpeg", "png", "bmp", "tiff"):
        return extract_text_from_image(file_path), links

    raise ValueError(f"Unsupported file type: {extension}")


def _clean_json(raw: str) -> str:
    if "```json" in raw:
        raw = raw.split("```json", 1)[1].split("```", 1)[0]
    elif "```" in raw:
        parts = raw.split("```")
        raw = parts[1] if len(parts) >= 2 else raw

    raw = raw.strip()
    if not raw.startswith("{"):
        start = raw.find("{")
        end = raw.rfind("}")
        if start != -1 and end != -1:
            raw = raw[start : end + 1]

    return raw


def _empty_cv_essentials() -> Dict:
    return {
        "full_name": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "summary": "",
        "skills": {"technical": [], "soft": []},
        "experience": [],
    }


def _empty_cv_full() -> Dict:
    return {
        "full_name": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "job_title": "",
        "summary": "",
        "education": [],
        "skills": {"technical": [], "soft": [], "tools": []},
        "experience": [],
        "certifications": [],
        "projects": [],
        "languages": [],
        "awards": [],
        "volunteer": [],
    }


def parse_cv_with_docling_llm(file_path: str) -> Dict:
    settings = get_recruitment_settings()
    llm_client = get_llm_client()

    text = extract_text_with_pymupdf(file_path).get("text", "")
    links = extract_links_from_pdf(file_path)

    if links["linkedin"]:
        text += "\n\n--- EXTRACTED LINKEDIN LINKS ---\n" + "\n".join(links["linkedin"])
    if links["github"]:
        text += "\n\n--- EXTRACTED GITHUB LINKS ---\n" + "\n".join(links["github"])

    docling_data = extract_with_docling(file_path)
    if docling_data and docling_data.get("text"):
        text += "\n\n--- DOCLING OUTPUT ---\n\n" + docling_data["text"]

    if not text:
        return _empty_cv_essentials()

    if len(text) > 10000:
        text = text[:10000]

    prompt = f"""You are an expert CV parser. Extract ONLY essential fields.

Return EXACTLY this JSON (no markdown, no preamble):
{{
  "full_name": "", "email": "", "phone": "",
  "linkedin": "full URL or empty", "github": "full URL or empty",
  "summary": "",
  "skills": {{"technical": [], "soft": []}},
  "experience": [
    {{"title": "", "company": "", "duration": "", "description": ""}}
  ]
}}

CV TEXT:
{text}"""

    for model in [settings.parsing_model, settings.parsing_fallback_model]:
        try:
            response = llm_client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "CV parser. Return only valid JSON."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=3000,
                temperature=0.1,
            )
            result = response.choices[0].message.content.strip()
            return json.loads(_clean_json(result))
        except Exception:
            continue

    return _empty_cv_essentials()


def parse_cv_with_llm(cv_text: str) -> Dict:
    settings = get_recruitment_settings()
    llm_client = get_llm_client()

    if len(cv_text) > 10000:
        cv_text = cv_text[:10000]

    prompt = f"""You are an expert CV parser. Extract ALL information.

Rules:
- LinkedIn / GitHub: extract the full URL, not just the label text.
- If a field is missing, use empty string or empty array.
- Infer job_title from experience/skills if not explicit.
- Do NOT include degrees in certifications.
- Volunteer work may be inferred from context.

Return EXACTLY this JSON (no markdown):
{{
  "full_name":"","email":"","phone":"","linkedin":"","github":"",
  "job_title":"","summary":"",
  "education":[{{"degree":"","institution":"","location":"","start_date":"","end_date":"","gpa":""}}],
  "skills":{{"technical":[],"soft":[],"tools":[]}},
  "experience":[{{"title":"","company":"","location":"","employment_type":"",
                  "start_date":"","end_date":"","duration":"",
                  "responsibilities":[],"achievements":[],"technologies":[]}}],
  "certifications":[{{"name":"","issuer":"","issue_date":""}}],
  "projects":[{{"name":"","description":"","technologies":[],"date":""}}],
  "languages":[{{"language":"","proficiency":""}}],
  "awards":[{{"title":"","issuer":"","description":""}}],
  "volunteer":[{{"role":"","organization":"","duration":"","description":""}}]
}}

CV TEXT:
{cv_text}"""

    try:
        response = llm_client.chat.completions.create(
            model=settings.parsing_model,
            messages=[
                {"role": "system", "content": "Professional CV parser. Return only valid JSON."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=4000,
            temperature=0.1,
        )
        result = response.choices[0].message.content.strip()
        return json.loads(_clean_json(result))
    except Exception:
        return _empty_cv_full()
